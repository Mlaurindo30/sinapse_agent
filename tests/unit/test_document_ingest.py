"""
Unit tests for scripts/document_ingest.py

Strategy:
- Patch `scripts.document_ingest.get_connection` to return an in-memory SQLite
  connection with the minimal schema, so tests never touch hive_mind.db.
- Use real python-docx and real bytes for PDF to exercise the actual extraction code.
"""

import json
import sqlite3
import sys
import hashlib
from pathlib import Path
from unittest.mock import patch, MagicMock

import pytest

# Ensure the project root is on sys.path so `from core.database import ...` works
_PROJECT_ROOT = Path(__file__).resolve().parents[2]
if str(_PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(_PROJECT_ROOT))

# A minimal valid PDF that has no real page content but is syntactically valid
MIN_PDF = b"""%PDF-1.0
1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj
2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj
3 0 obj<</Type/Page/MediaBox[0 0 3 3]>>endobj
xref
0 4
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
trailer<</Size 4/Root 1 0 R>>
startxref
190
%%EOF"""


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_in_memory_conn():
    """
    Return a MagicMock-wrapped in-memory SQLite connection.

    ingest_single_file calls conn.close() before returning, which would make
    the in-memory DB inaccessible to post-call assertions.  Wrapping in
    MagicMock lets us suppress close() while delegating every other method to
    the real sqlite3.Connection.
    """
    from unittest.mock import MagicMock
    conn = sqlite3.connect(":memory:")
    conn.row_factory = sqlite3.Row
    conn.executescript("""
        CREATE TABLE observations (
            id TEXT PRIMARY KEY,
            session_id TEXT,
            project TEXT,
            type TEXT,
            title TEXT,
            content TEXT,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            neuron_id TEXT,
            archived INTEGER DEFAULT 0,
            metadata JSON
        );

        CREATE TABLE document_memories (
            id TEXT PRIMARY KEY,
            file_path TEXT NOT NULL,
            file_hash TEXT UNIQUE,
            summary TEXT,
            topics TEXT,
            metadata JSON,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        );
    """)
    conn.commit()
    mock_conn = MagicMock(wraps=conn)
    mock_conn.close = MagicMock()  # suppress close so DB stays open for assertions
    return mock_conn, conn          # return both: mock for patching, real for querying


def _docx_available():
    try:
        from docx import Document  # noqa: F401
        return True
    except ImportError:
        return False


def _fitz_available():
    try:
        import fitz  # noqa: F401
        return True
    except ImportError:
        return False


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

class TestIngestDocx:
    """Tests for DOCX ingestion via ingest_single_file."""

    @pytest.mark.skipif(not _docx_available(), reason="python-docx not installed")
    def test_ingest_docx_creates_record(self, tmp_path):
        """should create a document_memories record when a valid DOCX is ingested"""
        from docx import Document as DocxDocument

        docx_path = tmp_path / "sample.docx"
        doc = DocxDocument()
        doc.add_paragraph("Hello Hive-Mind")
        doc.save(str(docx_path))

        mock_conn, real_conn = _make_in_memory_conn()

        with patch("scripts.document_ingest.get_connection", return_value=mock_conn):
            from scripts.document_ingest import ingest_single_file
            result = ingest_single_file(docx_path)

        assert result is True

        row = real_conn.execute(
            "SELECT * FROM document_memories WHERE file_path = ?",
            ("sample.docx",)
        ).fetchone()
        assert row is not None, "A record should exist in document_memories"

    @pytest.mark.skipif(not _docx_available(), reason="python-docx not installed")
    def test_ingest_docx_record_has_correct_hash(self, tmp_path):
        """should store the correct SHA-256 hash for the ingested DOCX"""
        from docx import Document as DocxDocument

        docx_path = tmp_path / "hash_check.docx"
        doc = DocxDocument()
        doc.add_paragraph("Content for hash check")
        doc.save(str(docx_path))

        expected_hash = hashlib.sha256(docx_path.read_bytes()).hexdigest()

        mock_conn, real_conn = _make_in_memory_conn()
        with patch("scripts.document_ingest.get_connection", return_value=mock_conn):
            from scripts.document_ingest import ingest_single_file
            ingest_single_file(docx_path)

        row = real_conn.execute(
            "SELECT file_hash FROM document_memories WHERE file_path = ?",
            ("hash_check.docx",)
        ).fetchone()
        assert row["file_hash"] == expected_hash


class TestIngestPdf:
    """Tests for PDF ingestion via ingest_single_file."""

    @pytest.mark.skipif(not _fitz_available(), reason="PyMuPDF not installed")
    def test_ingest_pdf_creates_record(self, tmp_path):
        """should create a document_memories record when a valid PDF is ingested"""
        pdf_path = tmp_path / "test.pdf"
        pdf_path.write_bytes(MIN_PDF)

        mock_conn, real_conn = _make_in_memory_conn()

        with patch("scripts.document_ingest.get_connection", return_value=mock_conn):
            from scripts.document_ingest import ingest_single_file
            result = ingest_single_file(pdf_path)

        assert result is True

        row = real_conn.execute(
            "SELECT * FROM document_memories WHERE file_path = ?",
            ("test.pdf",)
        ).fetchone()
        assert row is not None, "A record should exist in document_memories for the PDF"

    @pytest.mark.skipif(not _fitz_available(), reason="PyMuPDF not installed")
    def test_ingest_pdf_observation_is_created(self, tmp_path):
        """should create an observation record when a valid PDF is ingested"""
        pdf_path = tmp_path / "obs_test.pdf"
        pdf_path.write_bytes(MIN_PDF)

        mock_conn, real_conn = _make_in_memory_conn()

        with patch("scripts.document_ingest.get_connection", return_value=mock_conn):
            from scripts.document_ingest import ingest_single_file
            ingest_single_file(pdf_path)

        count = real_conn.execute(
            "SELECT COUNT(*) FROM observations WHERE type = 'document_ingest'"
        ).fetchone()[0]
        assert count == 1, "Exactly one observation should be created for the ingested PDF"


class TestIngestEdgeCases:
    """Edge case tests for ingest_single_file."""

    def test_ingest_nonexistent_returns_false(self):
        """should return False without raising an exception when the file does not exist"""
        from scripts.document_ingest import ingest_single_file

        result = ingest_single_file(Path("/nonexistent/path/file.docx"))
        assert result is False

    def test_ingest_unsupported_extension_returns_false(self, tmp_path):
        """should return False when given a file with an unsupported extension"""
        txt_path = tmp_path / "notes.txt"
        txt_path.write_text("Some text")

        from scripts.document_ingest import ingest_single_file
        result = ingest_single_file(txt_path)
        assert result is False

    @pytest.mark.skipif(not _docx_available(), reason="python-docx not installed")
    def test_ingest_idempotent(self, tmp_path):
        """should produce only one document_memories record when the same file is ingested twice"""
        from docx import Document as DocxDocument
        from scripts.document_ingest import ingest_single_file

        docx_path = tmp_path / "idempotent.docx"
        doc = DocxDocument()
        doc.add_paragraph("Idempotency test paragraph")
        doc.save(str(docx_path))

        mock_conn, real_conn = _make_in_memory_conn()

        # Both calls use the same mock — close() is suppressed so the DB stays open
        with patch("scripts.document_ingest.get_connection", return_value=mock_conn):
            ingest_single_file(docx_path)
            ingest_single_file(docx_path)

        count = real_conn.execute(
            "SELECT COUNT(*) FROM document_memories WHERE file_path = ?",
            ("idempotent.docx",)
        ).fetchone()[0]
        assert count == 1, "Ingesting the same file twice should not create duplicate records"

    @pytest.mark.skipif(not _docx_available(), reason="python-docx not installed")
    def test_ingest_idempotent_observation(self, tmp_path):
        """should not duplicate observations when the same file is ingested twice"""
        from docx import Document as DocxDocument
        from scripts.document_ingest import ingest_single_file

        docx_path = tmp_path / "idempotent_obs.docx"
        doc = DocxDocument()
        doc.add_paragraph("Observation idempotency test")
        doc.save(str(docx_path))

        mock_conn, real_conn = _make_in_memory_conn()

        with patch("scripts.document_ingest.get_connection", return_value=mock_conn):
            ingest_single_file(docx_path)
            ingest_single_file(docx_path)

        count = real_conn.execute(
            "SELECT COUNT(*) FROM observations WHERE title LIKE '%idempotent_obs.docx%'"
        ).fetchone()[0]
        assert count == 1, "INSERT OR IGNORE must prevent duplicate observations"
